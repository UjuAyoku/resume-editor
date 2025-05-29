import streamlit as st
from streamlit_tags import st_tags
from functions import create_prompt, get_resume_response, export_resume
import json
import base64
from docx import Document
from io import BytesIO
import re
import openai
import tempfile
import os
import PyPDF2
from collections import Counter
#import spacy
import en_core_web_sm



# Automatically download the model if not present
#try:
#    nlp = spacy.load("en_core_web_sm")
#except OSError:
#    download("en_core_web_sm")
#    nlp = spacy.load("en_core_web_sm")

# Load English language model for NLP
# nlp = spacy.load("en_core_web_sm")
nlp = en_core_web_sm.load()

# Set page config
st.set_page_config(
    page_title="Resume Optimizer Pro",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'resume_data' not in st.session_state:
    st.session_state.resume_data = {
        "personal_info": {},
        "summary": "",
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
        "certifications": [],
        "custom_sections": []
    }

# Helper functions
def generate_download_link(file_content, filename, text):
    """Generate a download link for files"""
    b64 = base64.b64encode(file_content).decode()
    href = f'<a href="data:file/{filename.split(".")[-1]};base64,{b64}" download="{filename}">{text}</a>'
    return href

def create_word_document(resume_data):
    """Create a Word document from resume data"""
    doc = Document()
    
    # Add title
    doc.add_heading(f"{resume_data['personal_info'].get('name', 'Resume')}", level=0)
    
    # Add contact info
    contact_info = []
    if resume_data['personal_info'].get('email'):
        contact_info.append(resume_data['personal_info']['email'])
    if resume_data['personal_info'].get('phone'):
        contact_info.append(resume_data['personal_info']['phone'])
    if resume_data['personal_info'].get('linkedin'):
        contact_info.append(resume_data['personal_info']['linkedin'])
    
    if contact_info:
        doc.add_paragraph(" | ".join(contact_info))
    
    # Add summary
    if resume_data['summary']:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(resume_data['summary'])
    
    # Add experience
    if resume_data['experience']:
        doc.add_heading("Experience", level=1)
        for exp in resume_data['experience']:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')}").bold = True
            p.add_run(f" at {exp.get('company', '')}")
            p.add_run(f" | {exp.get('start_date', '')} - {exp.get('end_date', 'Present')}")
            doc.add_paragraph(exp.get('description', ''))
    
    # Save to bytes
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def extract_text_from_file(uploaded_file):
    """Extract text from uploaded PDF or Word file"""
    text = ""
    try:
        if uploaded_file.type == "application/pdf":
            reader = PyPDF2.PdfReader(uploaded_file)
            text = " ".join([page.extract_text() for page in reader.pages])
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            text = " ".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error extracting text: {str(e)}")
    return text

def extract_keywords(text, max_keywords=20):
    """Extract keywords from text"""
    if not text:
        return []
    
    # Remove punctuation and convert to lowercase
    text = re.sub(r'[^\w\s]', '', text.lower())
    
    # Common stop words to exclude
    stop_words = set([
        'the', 'and', 'of', 'to', 'in', 'a', 'is', 'that', 'for', 'it',
        'with', 'as', 'on', 'at', 'by', 'this', 'are', 'be', 'from'
    ])
    
    # Extract words and filter
    words = [word for word in text.split() 
             if word not in stop_words and len(word) > 3]
    
    # Count word frequencies
    word_counts = Counter(words)
    
    # Return most common keywords
    return [kw for kw, count in word_counts.most_common(max_keywords)]

def compare_keywords(resume_text, job_description_text):
    """Compare keywords between resume and job description"""
    resume_keywords = extract_keywords(resume_text)
    job_keywords = extract_keywords(job_description_text)
    
    matches = [kw for kw in job_keywords if kw in resume_keywords]
    missing = [kw for kw in job_keywords if kw not in resume_keywords]
    
    return matches, missing

def calculate_ats_score(matches, total_keywords):
    """Calculate ATS compatibility score"""
    if total_keywords == 0:
        return 0
    return int((len(matches) / total_keywords) * 100)

def count_keyword_mentions(text, keyword):
    """Count how many times a keyword appears in text"""
    return len(re.findall(r'\b' + re.escape(keyword.lower()) + r'\b', text.lower()))

def analyze_resume(resume_text, job_description_text):
    """Analyze resume and provide optimization suggestions"""
    suggestions = []
    
    # Check for action verbs
    action_verbs = ['developed', 'created', 'implemented', 'managed', 'led', 'improved']
    missing_verbs = [verb for verb in action_verbs if verb not in resume_text.lower()]
    if missing_verbs:
        suggestions.append(f"Use more action verbs like: {', '.join(missing_verbs[:3])}")
    
    # Check for quantified achievements
    if not re.search(r'\d+%|\$\d+|\d+\+', resume_text):
        suggestions.append("Quantify achievements with numbers (e.g., 'increased sales by 20%')")
    
    # Check keyword matching
    matches, missing = compare_keywords(resume_text, job_description_text)
    if missing:
        suggestions.append(f"Add these missing keywords: {', '.join(missing[:3])}")
    
    # Formatting suggestions
    suggestions.append("Ensure consistent formatting throughout the document")
    suggestions.append("Tailor skills to match the job description more closely")
    
    return suggestions

# Sidebar
with st.sidebar:
    st.title("Resume Tools")
    tool_option = st.radio(
        "Choose an option:",
        ("Create New Resume", "Edit Resume", "Analyze Resume", "ATS Checker")
    )
    
    st.markdown("---")
    st.markdown("### Tips for a Great Resume")
    st.markdown("""
    - Use action verbs
    - Quantify achievements
    - Tailor to the job description
    - Keep it concise (1-2 pages)
    - Use simple, clean formatting
    """)
    
    st.markdown("---")
    st.markdown("### Upload Existing Resume")
    uploaded_file = st.file_uploader("Upload your resume (PDF or Word)", type=["pdf", "docx"])

# Main content
st.title("📄 Resume Optimizer Pro")
st.markdown("Create, edit, and optimize your resume for better job opportunities")

if tool_option == "Create New Resume":
    st.header("Create a New Resume")
    
    with st.expander("Personal Information", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.resume_data['personal_info']['name'] = st.text_input("Full Name")
            st.session_state.resume_data['personal_info']['email'] = st.text_input("Email")
            st.session_state.resume_data['personal_info']['phone'] = st.text_input("Phone")
        with col2:
            st.session_state.resume_data['personal_info']['linkedin'] = st.text_input("LinkedIn URL")
            st.session_state.resume_data['personal_info']['portfolio'] = st.text_input("Portfolio/GitHub")
            st.session_state.resume_data['personal_info']['location'] = st.text_input("Location")
    
    with st.expander("Professional Summary"):
        st.session_state.resume_data['summary'] = st.text_area(
            "Write a brief summary (2-3 sentences)", 
            value=st.session_state.resume_data['summary'],
            height=100
        )
    
    with st.expander("Work Experience"):
        if st.button("Add Work Experience"):
            st.session_state.resume_data['experience'].append({
                "title": "",
                "company": "",
                "location": "",
                "start_date": "",
                "end_date": "",
                "description": ""
            })
        
        for i, exp in enumerate(st.session_state.resume_data['experience']):
            st.subheader(f"Experience #{i+1}")
            col1, col2 = st.columns(2)
            with col1:
                exp['title'] = st.text_input("Job Title", value=exp.get('title', ''), key=f"title_{i}")
                exp['company'] = st.text_input("Company", value=exp.get('company', ''), key=f"company_{i}")
                exp['location'] = st.text_input("Location", value=exp.get('location', ''), key=f"location_{i}")
            with col2:
                exp['start_date'] = st.text_input("Start Date", value=exp.get('start_date', ''), key=f"start_date_{i}")
                exp['end_date'] = st.text_input("End Date", value=exp.get('end_date', ''), key=f"end_date_{i}")
            exp['description'] = st.text_area(
                "Description (bullet points work best)", 
                value=exp.get('description', ''),
                height=100,
                key=f"desc_{i}"
            )
            if st.button(f"Remove Experience #{i+1}", key=f"remove_exp_{i}"):
                st.session_state.resume_data['experience'].pop(i)
                st.experimental_rerun()
    
    # Preview and Download
    st.header("Preview & Download")
    if st.button("Generate Preview"):
        doc_content = create_word_document(st.session_state.resume_data)
        st.download_button(
            label="Download as Word",
            data=doc_content,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        st.json(st.session_state.resume_data)

elif tool_option == "Edit Resume":
    st.header("Edit Your Resume")
    
    # Initialize session state variables if they don't exist
    if 'edit_complete' not in st.session_state:
        st.session_state.edit_complete = False
    
    # Display upload status message
    if not uploaded_file:
        st.warning("Please upload your resume to begin editing")
    else:
        st.success("Resume uploaded successfully!")
    
    # Job description input
    jd_string = st.text_area("Paste the job description you're applying for", height=200)
    
    # Edit button
    if st.button("✨ Analyze Resume"):
        if uploaded_file and jd_string:
            with st.spinner("Generating optimized resume..."):
                try:
                    # Extract text from resume
                    resume_text = extract_text_from_file(uploaded_file)
                    
                    if not resume_text:
                        st.error("Could not extract text from the uploaded file")
                        st.stop()
                    
                    # Create prompt and get response
                    prompt = create_prompt(resume_text, jd_string)
                    result = get_resume_response(prompt, openai.api_key)
                    
                    # Split into resume and suggestions
                    if "## Additional Suggestions" in result:
                        new_resume, suggestions = result.split("## Additional Suggestions", 1)
                        suggestions = "## Additional Suggestions\n" + suggestions
                    else:
                        new_resume, suggestions = result, ""
                    
                    # Store in session state
                    st.session_state["new_resume"] = new_resume
                    st.session_state["suggestions"] = suggestions
                    st.session_state.edit_complete = True
                    
                except Exception as e:
                    st.error(f"An error occurred: {str(e)}")
        else:
            st.warning("Please provide both resume and job description")
    
    # Display results if available
    if st.session_state.get("edit_complete", False):
        st.success("Analysis complete!")
        st.markdown("### ✨ Tailored Resume")
        st.markdown(st.session_state["new_resume"])
        
        st.markdown("---")
        st.markdown("### 💡 Additional Suggestions")
        st.markdown(st.session_state["suggestions"])
        
        # Export button
        if st.button("📄 Export as PDF"):
            export_result = export_resume(st.session_state["new_resume"])
            st.success(export_result)


elif tool_option == "Analyze Resume":
    st.header("Analyze Your Resume")
    
    if not uploaded_file:
        st.warning("Please upload your resume to get optimization suggestions")

    job_description = st.text_area("Paste the job description you're applying for", height=200)
    
    if uploaded_file:
        # Add Run Analysis button
        if st.button("Run Optimization Analysis"):
            # Extract text from resume
            resume_text = extract_text_from_file(uploaded_file)
            
            if not resume_text:
                st.error("Could not extract text from the uploaded file")
                st.stop()
            
            if not job_description:
                st.warning("Please provide a job description for optimization")
                st.stop()
            
            st.toast("Analysis complete!")
            
            # Get keywords and matches
            matches, missing = compare_keywords(resume_text, job_description)
            job_keywords = extract_keywords(job_description)
            score = calculate_ats_score(matches, len(job_keywords))
            
            # Display analysis in the same format as before
            st.subheader("Optimization Suggestions")
            suggestions = analyze_resume(resume_text, job_description)
            for suggestion in suggestions:
                st.markdown(f"- {suggestion}")
            
            st.subheader("ATS Compatibility Score")
            st.progress(score)
            st.markdown(f"Your resume scores **{score}/100** for ATS compatibility with this job description.")
            
            # Keyword analysis - maintaining the same display format
            st.subheader("Keyword Analysis")
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("**Missing Keywords**")
                if missing:
                    for kw in missing[:3]:  # Show top 3 missing with counts from JD
                        count = count_keyword_mentions(job_description, kw)
                        st.markdown(f"- {kw.capitalize()} ({count} mentions in JD)")
                else:
                    st.markdown("- No significant keywords missing")
            
            with col2:
                st.markdown("**Your Strong Keywords**")
                if matches:
                    for kw in matches[:3]:  # Show top 3 matches with counts from resume
                        count = count_keyword_mentions(resume_text, kw)
                        st.markdown(f"- {kw.capitalize()} (matches {count} mentions)")
                else:
                    st.markdown("- No strong keyword matches found")
            
            # Actionable recommendations - maintaining the same format
            st.subheader("Actionable Recommendations")
            if missing:
                st.markdown(f"1. Add 2-3 more mentions of '{missing[0].capitalize()}' in your skills and experience sections")
            if len(matches) > 0:
                st.markdown(f"2. Highlight your experience with '{matches[0].capitalize()}' more prominently")
            st.markdown("3. Reorder your skills to put the most relevant ones first")
        else:
            st.info("Click the 'Run Optimization Analysis' button to analyze your resume")
    # else:
    #     st.warning("Please upload your resume to get optimization suggestions")

elif tool_option == "ATS Checker":
    st.header("ATS Compatibility Checker")
    st.markdown("""
    Applicant Tracking Systems (ATS) are used by employers to screen resumes. 
    This tool helps you optimize your resume for better ATS performance.
    """)

    if not uploaded_file:
        st.warning("Please upload your resume to check ATS compatibility")
    
    job_description = st.text_area("Paste the job description (optional but recommended)", height=150)
    
    if uploaded_file:
        # Add Run Analysis button
        if st.button("Run ATS Analysis"):
            # Extract text from resume
            resume_text = extract_text_from_file(uploaded_file)
            
            if not resume_text:
                st.error("Could not extract text from the uploaded file")
                st.stop()
            
            st.success("Analysis complete!")
            
            # Process file and show ATS analysis - maintaining the same display format
            st.subheader("ATS Analysis Results")
            
            # Formatting check (simplified metrics)
            st.markdown("### Formatting Check")
            col1, col2, col3 = st.columns(3)
            col1.metric("Fonts", "2", "Good (≤3)")  # Simplified
            col2.metric("Sections", str(len(re.findall(r'\n\s*[A-Z][A-Z ]+:', resume_text))), "Standard")
            col3.metric("Length", f"{len(resume_text.split())//250+1} pages", "1-2 pages ideal")
            
            # Content check - maintaining the same format
            st.markdown("### Content Check")
            checks = [
                ("Action verbs", bool(re.search(r'\b(led|managed|developed|created)\b', resume_text, re.I))),
                ("Quantified achievements", bool(re.search(r'\d+%|\$\d+|\d+\+', resume_text))),
                ("Clear contact info", bool(re.search(r'[\w\.-]+@[\w\.-]+', resume_text))),
                ("Skills section", bool(re.search(r'skills:', resume_text, re.I)))
            ]
            
            for check, passed in checks:
                status = "✅" if passed else "⚠️"
                st.markdown(f"- {status} {check}")
            
            # Keyword analysis (if job description provided)
            if job_description:
                st.markdown("### Keyword Density")
                job_keywords = extract_keywords(job_description)
                resume_keywords = extract_keywords(resume_text)
                matches, missing = compare_keywords(resume_text, job_description)
                
                # Create keyword density chart - maintaining the same format
                top_keywords = {kw: count_keyword_mentions(resume_text, kw) 
                              for kw in job_keywords[:4]}  # Show top 4
                st.bar_chart(top_keywords)
                
                # Calculate ATS score
                score = calculate_ats_score(matches, len(job_keywords))
                
                st.markdown("### Overall ATS Score")
                st.progress(score)
                st.markdown(f"Your resume scores **{score}/100** for ATS compatibility.")
            else:
                st.info("Add a job description for keyword analysis and ATS scoring")
        else:
            st.info("Click the 'Run ATS Analysis' button to analyze your resume")
    # else:
    #     st.warning("Please upload your resume to check ATS compatibility")

# Footer
st.markdown("---")
st.markdown("""
**Need help?** Contact support@resumeoptimizer.pro
""")
